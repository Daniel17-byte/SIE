from pathlib import Path
import re
import unicodedata

from docx import Document
from docx.shared import Pt, RGBColor

BASE = Path('/Users/daniellungu/Desktop/UT AC/SIE/Curs/SIE/materiale')
PARTIAL_PATH = BASE / 'Partial.txt'
EXAMEN_PATH = BASE / 'Examen.txt'
OUT_PATH = BASE / 'Intrebari_Partial_Examen.docx'
TITLE = 'Intrebari grupate pe teme similare (Partial / Examen)'

# Matches both "Intrebarea 12" and "15...." styles.
ENTRY_START_RE = re.compile(r'^(?:(?:[Ii\u00ce\u00ee])ntrebarea\s+\d+|\d+\.)')
QUESTION_NUMBER_ONLY_RE = re.compile(r'^(?:[Ii\u00ce\u00ee])ntrebarea\s+\d+\s*$')
WORD_RE = re.compile(r"[A-Za-z\u0102\u00c2\u00ce\u0218\u021a\u0103\u00e2\u00ee\u0219\u021b0-9\-]+")
CHAPTER_RE = re.compile(r'^📂\s*(.+)$')
ORIGIN_RE = re.compile(r'^Capitol original:\s*(.+)$')

STOPWORDS = {
    'si', 'sau', 'de', 'la', 'cu', 'din', 'pe', 'pentru', 'care', 'ale', 'al', 'a', 'ai',
    'ce', 'cum', 'in', 'fata', 'versus', 'vs', 'prin', 'utilizate', 'utilizata', 'utilizat',
    'utilizand', 'utilizare', 'explicati', 'prezentati', 'descrieti', 'diferenta',
    'tehnologia', 'tehnicile', 'principiul', 'structura', 'avantajele', 'dezavantajele',
    'caracteristicile', 'comparati'
}

CHAPTER_ORDER = [
    'Introducere',
    'Metode IE',
    'Magistrale',
    'Module Extensie',
    'Afisaje',
    'Adaptoare Grafice',
    'Discuri Optice',
    'Diverse',
]

SECTION_ORDER = ['Partial', 'Examen']

THEME_RULES = {
    'Partial': [
        ('Diagrame de timp si protocoale de transfer', ['diagrama', 'protocol', 'transfer initiat', 'operatie de citire', 'operatie de scriere', 'bidirectional', 'unidirectional']),
        ('USB', ['usb4', 'usb 3.0', 'usb 3.1', 'usb 3.2', 'magistralei usb', 'transferurile izocrone', 'transferurile asincrone pe usb']),
        ('PCI / PCI Express', ['pci express', 'pcie', 'magistrala pci']),
        ('Arbitraj si controlul magistralei', ['arbitraj', 'arbitrarea', 'daisy', 'busreq', 'busgnt', 'busy']),
        ('DMA si procesoare de I/E', ['dma', 'furt de ciclu', 'rafala', 'pie', 'procesorul de i/e', 'procesoarelor de ie']),
        ('Intreruperi', ['intreruper', 'iack', 'rutinei de tratare']),
        ('Formate si module industriale', ['vxs', 'compactpci', 'plusio', 'com express', 'mezanin']),
        ('Magistrale seriale/paralele si interfete', ['seriala', 'seriale', 'paralela', 'paralele', 'i2c', 'spi', 'vme320', 'magistrala locala']),
        ('Bazele I/E si arhitectura sistemului', ['timpul de executie', 'modul de i/e', 'operatie de i/e programata', 'ie programata', 'ie programate', 'interogare hardware', 'embedded', 'periferice', 'mapare in memorie', 'adresarea izolata']),
    ],
    'Examen': [
        ('LCD - fundamente si structura', ['nematica', 'colesterica', 'adresarea directa', 'adresarea multiplexata', 'matrice pasiva', 'matrice activa', 'pixeli defectuosi', 'cristale lichide']),
        ('IPS / H-IPS / S-IPS', ['ips', 'h-ips', 's-ips', 'super ips', 'in-plane switching']),
        ('VA / MVA / PVA', ['vertical alignment', 'mva', 'pva', 's-pva', 'aliniere verticala']),
        ('Iluminare de fundal LCD', ['ccfl', 'luminii de fond', 'iluminarii de fundal', 'led rgb', 'fald', 'mini-diode led', 'variantele de iluminare de fundal']),
        ('Culori, contrast, quantum dots si timp de raspuns', ['rtc', 'frame rate control', 'spatial dithering', 'contrast', 'luminozitate', 'cromaticitate', 'saturatie', 'gama de culori', 'qdef', 'quantum dot', 'puncte cuantice', 'foto-emisive', 'unghiurilor de vizualizare']),
        ('OLED', ['oled', 'amoled', 'pmoled', 'woled', 'foled', 'diamond pixel', 'inkjet printing', 'fluorescente', 'fosforescente', 'sub-pixel', 'sub-pixelilor']),
        ('E-paper', ['hartie electronica', 'e-paper', 'electroforetica', 'micro-capsule', 'modulare interferometrica', 'cerneala cu trei pigmenti']),
        ('Interfete video si codificari', ['displayport', 'hdmi', 'nrzi']),
        ('Memorie grafica', ['gddr6', 'high bandwidth memory', 'hbm', 'memoria grafica']),
        ('GPU, CUDA si arhitecturi grafice', ['cuda', 'gpu', 'ucp', 'adaptor grafic', 'kernel']),
        ('Pipeline grafic si texturare', ['pipeline', 'rasterizare', 'triunghiului', 'filtrare a texturilor', 'tesselare', 'procesare geometrica']),
        ('Discuri optice', ['truex', 'cd-rw', 'dvd', 'blu-ray', 'ansamblu optic', 'discurile compact', 'cd-r', 'canelurii spiralate', 'clv']),
    ],
}


def split_entries(path: Path) -> list[list[str]]:
    entries: list[list[str]] = []
    current: list[str] = []

    for raw in path.read_text(encoding='utf-8').splitlines():
        line = raw.rstrip()
        if ENTRY_START_RE.match(line.strip()):
            if current:
                entries.append(current)
            current = [line]
        elif current:
            current.append(line)

    if current:
        entries.append(current)

    return entries


def parse_entry(entry_lines: list[str]) -> tuple[str, str, list[str]]:
    chapter = ''
    cleaned: list[str] = []

    for raw in entry_lines:
        line = raw.strip()
        if not line:
            continue
        chapter_match = CHAPTER_RE.match(line)
        if chapter_match:
            chapter = chapter_match.group(1).strip()
            continue
        if line.startswith('📖 Sursă:'):
            continue
        cleaned.append(line)

    if not cleaned:
        return ('', chapter, [])

    head = cleaned[0]
    if QUESTION_NUMBER_ONLY_RE.match(head):
        if len(cleaned) == 1:
            return ('', chapter, [])
        question = cleaned[1]
        answer = cleaned[2:]
        return (question, chapter, answer)

    question = re.sub(r'^\d+\.\s*', '', head).strip()
    answer = cleaned[1:]
    return (question, chapter, answer)


def extract_qas(path: Path) -> list[dict]:
    entries = split_entries(path)
    qas: list[dict] = []

    for entry in entries:
        question, chapter, answer = parse_entry(entry)
        if question:
            qas.append({'question': question, 'chapter': chapter, 'answer': answer})

    return qas


def read_docx_sections(path: Path) -> dict[str, list[dict]]:
    doc = Document(path)
    sections = {section: [] for section in SECTION_ORDER}
    current_section = ''
    current_group = ''
    current_qa = None

    def flush_current() -> None:
        nonlocal current_qa
        if current_section and current_qa:
            sections[current_section].append(current_qa)
            current_qa = None

    for paragraph in doc.paragraphs:
        raw = paragraph.text or ''
        text = raw.strip()
        if not text:
            continue
        if text in {TITLE, 'Intrebari - Partial apoi Examen'}:
            continue
        if text in SECTION_ORDER:
            flush_current()
            current_section = text
            current_group = ''
            continue
        if text.startswith('Capitol: ') or text.startswith('Tema: '):
            flush_current()
            current_group = text.split(':', 1)[1].strip()
            continue
        if raw.startswith('\t'):
            if not current_qa:
                continue
            origin_match = ORIGIN_RE.match(text)
            if origin_match:
                current_qa['chapter'] = origin_match.group(1).strip()
            else:
                current_qa['answer'].append(text)
            continue

        flush_current()
        current_qa = {
            'question': text,
            'chapter': current_group if current_group in CHAPTER_ORDER else '',
            'answer': [],
        }

    flush_current()
    return sections


def fold_text(value: str) -> str:
    # Remove accents so fuzzy matching works for both diacritics and non-diacritics text.
    return ''.join(c for c in unicodedata.normalize('NFKD', value.lower()) if not unicodedata.combining(c))


def compact_spaces(value: str) -> str:
    return re.sub(r'\s+', ' ', value).strip()


def question_token_set(question: str) -> set[str]:
    plain = fold_text(re.sub(r'^(?:(?:[Ii\u00ce\u00ee])ntrebarea\s+\d+\s*|\d+\.)\s*', '', question))
    tokens: set[str] = set()
    for token in re.findall(r'[a-z0-9\-]+', plain):
        if len(token) < 4 or token in STOPWORDS:
            continue
        tokens.add(token)
    return tokens


def has_conflicting_terms(a: str, b: str) -> bool:
    af = fold_text(a)
    bf = fold_text(b)
    conflicts = [
        ('avantaje', 'dezavantaje'),
        ('citire', 'scriere'),
        ('sursa', 'destinatie'),
        ('sursa', 'destinatia'),
        ('directa', 'multiplexata'),
    ]

    for left, right in conflicts:
        if (left in af and right in bf) or (right in af and left in bf):
            return True
    return False


def is_similar_question(a: str, b: str) -> bool:
    if has_conflicting_terms(a, b):
        return False

    af = fold_text(re.sub(r'\s+', ' ', a).strip())
    bf = fold_text(re.sub(r'\s+', ' ', b).strip())
    if af == bf:
        return True

    ta = question_token_set(a)
    tb = question_token_set(b)
    if not ta or not tb:
        return False
    inter = len(ta & tb)
    union = len(ta | tb)
    if union == 0:
        return False
    score = inter / union
    # Merge questions that are ~90% the same (near-identical formulations).
    return score >= 0.90


def dedupe_similar_qas(qas: list[dict]) -> tuple[list[dict], list[tuple[str, str]]]:
    kept: list[dict] = []
    removed: list[tuple[str, str]] = []

    for qa in qas:
        duplicate_of = None
        for existing in kept:
            if is_similar_question(qa['question'], existing['question']):
                duplicate_of = existing['question']
                break
        if duplicate_of:
            removed.append((qa['question'], duplicate_of))
            continue
        kept.append(qa)

    return kept, removed


def dedupe_exact_questions(qas: list[dict]) -> tuple[list[dict], list[tuple[str, str]]]:
    kept: list[dict] = []
    seen: dict[str, str] = {}
    removed: list[tuple[str, str]] = []

    for qa in qas:
        normalized = fold_text(compact_spaces(qa['question']))
        if normalized in seen:
            removed.append((qa['question'], seen[normalized]))
            continue
        seen[normalized] = qa['question']
        kept.append(qa)

    return kept, removed


def classify_theme(section_name: str, qa: dict) -> str:
    haystack = fold_text(f"{qa['question']} {qa.get('chapter', '')}")
    for theme_name, keywords in THEME_RULES[section_name]:
        for keyword in keywords:
            if fold_text(keyword) in haystack:
                return theme_name
    chapter = qa.get('chapter', '').strip()
    return chapter if chapter else 'Diverse'


def group_by_theme(section_name: str, qas: list[dict]) -> dict[str, list[dict]]:
    grouped: dict[str, list[dict]] = {}
    for qa in qas:
        theme = classify_theme(section_name, qa)
        grouped.setdefault(theme, []).append(qa)
    return grouped


def pick_keywords(question: str, max_keywords: int = 3) -> list[str]:
    q = re.sub(r'^(?:(?:[Ii\u00ce\u00ee])ntrebarea\s+\d+\s*|\d+\.)\s*', '', question)

    candidates: list[str] = []
    for match in WORD_RE.finditer(q):
        word = match.group(0)
        low = word.lower()
        if len(low) < 4 or low in STOPWORDS or low.isdigit():
            continue
        candidates.append(word)

    seen: set[str] = set()
    unique: list[str] = []
    for word in candidates:
        low = word.lower()
        if low not in seen:
            seen.add(low)
            unique.append(word)

    ranked = sorted(enumerate(unique), key=lambda x: (-len(x[1]), x[0]))[:max_keywords]
    return [word for _, word in sorted(ranked, key=lambda x: x[0])]


def pick_answer_keywords(answer_lines: list[str], max_keywords: int = 5) -> set[str]:
    text = ' '.join(answer_lines)
    candidates: list[str] = []
    for match in WORD_RE.finditer(text):
        word = match.group(0)
        low = fold_text(word)
        if len(low) < 4 or low in STOPWORDS or low.isdigit():
            continue
        candidates.append(word)

    seen: set[str] = set()
    unique: list[str] = []
    for word in candidates:
        low = fold_text(word)
        if low not in seen:
            seen.add(low)
            unique.append(word)

    ranked = sorted(enumerate(unique), key=lambda x: (-len(x[1]), x[0]))[:max_keywords]
    return {fold_text(word) for _, word in sorted(ranked, key=lambda x: x[0])}


def add_highlighted_text(paragraph, text: str, keywords: set[str], *, size: int, bold: bool = False) -> None:
    for token in re.split(r'(\s+)', text):
        run = paragraph.add_run(token)
        style_run(run, size, bold=bold)

        clean_word = re.sub(
            r'^[^A-Za-z\u0102\u00c2\u00ce\u0218\u021a\u0103\u00e2\u00ee\u0219\u021b0-9\-]+|'
            r'[^A-Za-z\u0102\u00c2\u00ce\u0218\u021a\u0103\u00e2\u00ee\u0219\u021b0-9\-]+$',
            '',
            token,
        )
        if clean_word and fold_text(clean_word) in keywords:
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


def style_run(run, size: int, bold: bool = False, italic: bool = False, color=None) -> None:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_section(doc: Document, section_name: str, qas: list[dict]) -> None:
    header = doc.add_paragraph().add_run(section_name)
    style_run(header, 16, bold=True)

    grouped = group_by_theme(section_name, qas)
    theme_order = [theme for theme, _ in THEME_RULES[section_name] if theme in grouped]
    theme_order.extend(sorted(theme for theme in grouped if theme not in theme_order))

    for theme in theme_order:
        theme_header = doc.add_paragraph().add_run(f'Tema: {theme}')
        style_run(theme_header, 14, bold=True)

        for qa in grouped[theme]:
            question = qa['question']
            answer_lines = qa['answer']
            p = doc.add_paragraph(style='List Number')
            add_highlighted_text(p, question, set(), size=14, bold=True)

            chapter = qa.get('chapter', '').strip()
            if chapter:
                meta_p = doc.add_paragraph('\tCapitol original: ' + chapter)
                for run in meta_p.runs:
                    style_run(run, 11, italic=True, color=RGBColor(0x66, 0x66, 0x66))

            answer_keywords = pick_answer_keywords(answer_lines)
            for ans in answer_lines:
                # Keep answer lines visually nested under each question.
                ans_p = doc.add_paragraph('\t')
                add_highlighted_text(ans_p, ans, answer_keywords, size=12)


def load_source_sections() -> dict[str, list[dict]]:
    if OUT_PATH.exists():
        return read_docx_sections(OUT_PATH)

    partial_qas = extract_qas(PARTIAL_PATH)
    examen_qas = extract_qas(EXAMEN_PATH)
    return {'Partial': partial_qas, 'Examen': examen_qas}


def main() -> None:
    source_sections = load_source_sections()

    partial_qas_raw = source_sections.get('Partial', [])
    examen_qas_raw = source_sections.get('Examen', [])

    partial_qas, partial_removed_exact = dedupe_exact_questions(partial_qas_raw)
    examen_qas, examen_removed_exact = dedupe_exact_questions(examen_qas_raw)

    partial_qas, partial_removed_similar = dedupe_similar_qas(partial_qas)
    examen_qas, examen_removed_similar = dedupe_similar_qas(examen_qas)

    partial_removed = partial_removed_exact + partial_removed_similar
    examen_removed = examen_removed_exact + examen_removed_similar

    doc = Document()
    normal = doc.styles['Normal'].font
    normal.name = 'Times New Roman'
    normal.size = Pt(12)

    title = doc.add_paragraph().add_run(TITLE)
    style_run(title, 18, bold=True)

    add_section(doc, 'Partial', partial_qas)
    add_section(doc, 'Examen', examen_qas)

    doc.save(OUT_PATH)

    print(f'Created: {OUT_PATH}')
    print(f'Partial entries: {len(partial_qas)} (removed duplicate titles: {len(partial_removed)})')
    print(f'Examen entries: {len(examen_qas)} (removed duplicate titles: {len(examen_removed)})')


if __name__ == '__main__':
    main()

